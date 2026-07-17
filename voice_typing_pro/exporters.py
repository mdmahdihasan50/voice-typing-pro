from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from PySide6.QtGui import QFont, QTextDocument
from PySide6.QtPrintSupport import QPrinter


def export_docx(text: str, path: str | Path, font_family: str, font_size: int) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = font_family
    style.font.size = Pt(font_size)
    for block in text.splitlines() or [""]:
        document.add_paragraph(block)
    document.save(str(path))


def export_pdf(text: str, path: str | Path, font_family: str, font_size: int) -> None:
    printer = QPrinter(QPrinter.PrinterMode.HighResolution)
    printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
    printer.setOutputFileName(str(path))
    document = QTextDocument()
    document.setDefaultFont(QFont(font_family, font_size))
    document.setPlainText(text)
    document.print_(printer)
